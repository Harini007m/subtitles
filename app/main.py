import os
import re
import concurrent.futures
from fastapi import FastAPI, UploadFile, File, Form, Body
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse
from fastapi.templating import Jinja2Templates
from fastapi.requests import Request
from fastapi.staticfiles import StaticFiles

from app.transcription import transcribe_audio
from app.video_utils import extract_audio, embed_subtitles, embed_subtitles_mp4, remux_to_mp4

app = FastAPI()

app.mount("/static", StaticFiles(directory="static"), name="static")
templates = Jinja2Templates(directory="templates")

UPLOAD_DIR = "uploads"
OUTPUT_DIR = "outputs"

os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)

# In-memory cache: safe_filename -> original segments (English)
transcription_cache = {}

# Cache for the latest burned-in video per filename
burnin_cache = {}


def safe_filename(name: str) -> str:
    """Remove special chars and spaces from filename to avoid ffmpeg/path issues."""
    base, ext = os.path.splitext(name)
    base = re.sub(r"[^\w\-]", "_", base)   # keep letters, digits, _, -
    return base + ext.lower()


def segments_to_srt(segments: list) -> str:
    """Convert segment list to SRT format string."""
    def fmt(sec):
        h = int(sec // 3600)
        m = int((sec % 3600) // 60)
        s = sec % 60
        ms = int((s - int(s)) * 1000)
        return f"{h:02d}:{m:02d}:{int(s):02d},{ms:03d}"

    lines = []
    for i, seg in enumerate(segments, start=1):
        lines.append(str(i))
        lines.append(f"{fmt(seg['start'])} --> {fmt(seg['end'])}")
        lines.append(seg["text"].strip())
        lines.append("")
    return "\n".join(lines)


@app.get("/", response_class=HTMLResponse)
async def home(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})


@app.post("/transcribe/")
async def transcribe_video(file: UploadFile = File(...)):
    """
    Step 1: Upload video and transcribe it. Returns JSON segments.
    """
    fname = safe_filename(file.filename)
    video_path = os.path.join(UPLOAD_DIR, fname)

    # Write uploaded file to disk
    with open(video_path, "wb") as f:
        content = await file.read()
        f.write(content)

    try:
        # Extract audio (works for any video format)
        audio_path = extract_audio(video_path)

        # Transcribe with Whisper
        segments = transcribe_audio(audio_path)
    except Exception as e:
        print(f"[transcription error] {e}")
        return JSONResponse(
            {"error": f"Transcription failed: {str(e)}"},
            status_code=500
        )

    # Cache original segments
    transcription_cache[fname] = segments

    return JSONResponse({
        "filename": fname,
        "segments": segments
    })


@app.get("/video/{filename}")
async def stream_video(filename: str):
    """Serve the uploaded video for playback."""
    video_path = os.path.join(UPLOAD_DIR, filename)
    if not os.path.exists(video_path):
        return JSONResponse({"error": "File not found"}, status_code=404)

    # Pick correct MIME type
    ext = os.path.splitext(filename)[1].lower()
    mime_map = {
        ".mp4": "video/mp4",
        ".mkv": "video/x-matroska",
        ".mov": "video/quicktime",
        ".avi": "video/x-msvideo",
        ".webm": "video/webm",
    }
    media_type = mime_map.get(ext, "video/mp4")

    return FileResponse(video_path, media_type=media_type)


@app.post("/translate/")
async def translate_segments(
    filename: str = Form(...),
    language: str = Form(...),
):
    """
    Step 2 (live): Translate already-transcribed segments into a new language.
    Returns translated JSON segments instantly (no video re-encoding needed).
    """
    if filename not in transcription_cache:
        return JSONResponse(
            {"error": "Video not transcribed yet. Please upload again."},
            status_code=400
        )

    segments = transcription_cache[filename]
    all_texts = [seg["text"] for seg in segments]

    if language == "en":
        translated_texts = all_texts
    else:
        try:
            from deep_translator import GoogleTranslator
            translator = GoogleTranslator(source="auto", target=language)

            def translate_single(text):
                try:
                    result = translator.translate(text)
                    return result if result else text
                except Exception:
                    return text

            with concurrent.futures.ThreadPoolExecutor(max_workers=10) as executor:
                translated_texts = list(executor.map(translate_single, all_texts))

        except Exception as e:
            print(f"[translation error] {e}")
            translated_texts = all_texts

    translated_segments = [
        {
            "start": seg["start"],
            "end": seg["end"],
            "text": translated_texts[i]
        }
        for i, seg in enumerate(segments)
    ]

    return JSONResponse({"segments": translated_segments})


@app.post("/burnin/")
async def burnin_subtitles(payload: dict = Body(...)):
    """
    Author-only: Accept edited segments + filename, burn subtitles into video,
    cache the output path, and return the output filename (always .mp4).
    """
    filename = payload.get("filename")
    segments = payload.get("segments", [])

    if not filename:
        return JSONResponse({"error": "filename is required"}, status_code=400)

    video_path = os.path.join(UPLOAD_DIR, filename)
    if not os.path.exists(video_path):
        return JSONResponse({"error": "Original video not found"}, status_code=404)

    # Write SRT to outputs dir
    srt_filename = os.path.splitext(filename)[0] + "_subtitle.srt"
    srt_path = os.path.join(OUTPUT_DIR, srt_filename)
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    srt_content = segments_to_srt(segments)
    with open(srt_path, "w", encoding="utf-8") as f:
        f.write(srt_content)

    try:
        output_video_path = embed_subtitles_mp4(video_path, srt_path)
    except Exception as e:
        print(f"[burnin error] {e}")
        return JSONResponse({"error": f"Burn-in failed: {str(e)}"}, status_code=500)

    output_basename = os.path.basename(output_video_path)
    burnin_cache[filename] = output_video_path

    return JSONResponse({"output_filename": output_basename})


@app.get("/download-video/{filename}")
async def download_output_video(filename: str):
    """Download the burned-in subtitle video."""
    video_path = os.path.join(OUTPUT_DIR, filename)
    if not os.path.exists(video_path):
        return JSONResponse({"error": "Output video not found"}, status_code=404)

    return FileResponse(
        video_path,
        media_type="video/mp4",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )


@app.get("/download-mp4/{filename}")
async def download_original_as_mp4(filename: str):
    """
    Download the original uploaded video as an MP4 file.
    If the upload is already .mp4, serve it directly.
    Otherwise re-mux it to MP4 with ffmpeg (fast, no re-encode) and cache the result.
    """
    upload_path = os.path.join(UPLOAD_DIR, filename)
    if not os.path.exists(upload_path):
        return JSONResponse({"error": "Video not found"}, status_code=404)

    ext = os.path.splitext(filename)[1].lower()
    if ext == ".mp4":
        # Already MP4 — serve directly
        dl_name = filename
        serve_path = upload_path
    else:
        # Re-mux (copy streams) to MP4 — very fast, no quality loss
        base_name = os.path.splitext(filename)[0]
        dl_name = base_name + ".mp4"
        serve_path = os.path.join(OUTPUT_DIR, f"original_{dl_name}")
        if not os.path.exists(serve_path):
            try:
                remux_to_mp4(upload_path, serve_path)
            except Exception as e:
                print(f"[remux error] {e}")
                return JSONResponse({"error": f"Conversion failed: {str(e)}"}, status_code=500)

    return FileResponse(
        serve_path,
        media_type="video/mp4",
        headers={"Content-Disposition": f'attachment; filename="{dl_name}"'}
    )


@app.post("/download-transcript/")
async def download_transcript(payload: dict = Body(...)):
    """
    Return a Word (.docx) transcript of the current segments.
    """
    from docx import Document
    from docx.shared import Pt

    filename = payload.get("filename")
    segments = payload.get("segments", [])

    if not filename:
        return JSONResponse({"error": "filename is required"}, status_code=400)

    def fmt_time(sec):
        h = int(sec // 3600)
        m = int((sec % 3600) // 60)
        s = int(sec % 60)
        return f"{h:02d}:{m:02d}:{s:02d}"

    doc = Document()

    # ── Title (plain paragraph + run — avoids heading.runs[] IndexError) ──
    title_para = doc.add_paragraph()
    title_para.alignment = 1          # 1 = CENTER
    title_run = title_para.add_run("Transcript")
    title_run.bold = True
    title_run.font.size = Pt(20)

    # ── Source filename ──
    src_para = doc.add_paragraph()
    src_para.alignment = 1            # CENTER
    src_run = src_para.add_run(
        "Source: " + os.path.splitext(filename)[0].replace("_", " ")
    )
    src_run.font.size = Pt(10)

    doc.add_paragraph()               # blank spacer line

    # ── One paragraph per subtitle segment ──
    for seg in segments:
        ts = f"[{fmt_time(seg['start'])}  ->  {fmt_time(seg['end'])}]"
        text = seg["text"].strip()

        para = doc.add_paragraph()
        # Timestamp run — bold, small
        ts_run = para.add_run(ts + "   ")
        ts_run.bold = True
        ts_run.font.size = Pt(9)
        # Subtitle text run
        txt_run = para.add_run(text)
        txt_run.font.size = Pt(11)
        para.paragraph_format.space_after = Pt(4)

    base = os.path.splitext(filename)[0]
    docx_filename = base + "_transcript.docx"
    docx_path = os.path.join(OUTPUT_DIR, docx_filename)
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    doc.save(docx_path)

    return FileResponse(
        docx_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{docx_filename}"'}
    )
